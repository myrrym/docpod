"""Parsing: extension dispatch, txt/md passthrough, docx and PDF extraction,
scanned-PDF detection, empty-document rejection."""

import pytest

from docpod.errors import EmptyDocumentError, ScannedPDFError, UnsupportedFormatError
from docpod.parsing import parse_document

LOREM = (
    "Local-first tooling keeps user data on the user's machine. "
    "This paragraph exists to give the parser something substantial to extract, "
    "well past any scanned-page character thresholds."
)


def test_txt_passthrough(tmp_path):
    f = tmp_path / "notes.txt"
    f.write_text(LOREM, encoding="utf-8")
    assert parse_document(f) == LOREM


def test_md_passthrough(tmp_path):
    f = tmp_path / "notes.md"
    f.write_text("# Title\n\n" + LOREM, encoding="utf-8")
    assert LOREM in parse_document(f)


def test_unsupported_extension(tmp_path):
    f = tmp_path / "slides.pptx"
    f.write_bytes(b"\x00")
    with pytest.raises(UnsupportedFormatError, match="pptx"):
        parse_document(f)


def test_empty_document(tmp_path):
    f = tmp_path / "empty.txt"
    f.write_text("   \n\n  ", encoding="utf-8")
    with pytest.raises(EmptyDocumentError):
        parse_document(f)


def test_docx(tmp_path):
    import docx

    f = tmp_path / "report.docx"
    d = docx.Document()
    d.add_paragraph("The first paragraph of the report.")
    d.add_paragraph(LOREM)
    d.save(str(f))

    text = parse_document(f)
    assert "The first paragraph of the report." in text
    assert LOREM in text


def test_pdf(tmp_path):
    import fitz

    f = tmp_path / "paper.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), LOREM, fontsize=11)
    doc.save(str(f))
    doc.close()

    assert "Local-first tooling" in parse_document(f)


def test_scanned_pdf_fails_clearly(tmp_path):
    import fitz

    f = tmp_path / "scan.pdf"
    doc = fitz.open()
    doc.new_page()  # image-only page: no text layer at all
    doc.new_page()
    doc.save(str(f))
    doc.close()

    with pytest.raises(ScannedPDFError, match="OCR"):
        parse_document(f)
