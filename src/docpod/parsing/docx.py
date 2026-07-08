from __future__ import annotations

from pathlib import Path


def parse_docx(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    # Tables are common in .docx reports; flatten rows as tab-separated lines.
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n\n".join(p for p in parts if p.strip())
